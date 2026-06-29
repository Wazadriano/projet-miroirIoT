# -*- coding: utf-8 -*-
"""
Genere le dossier Word de soutenance RNCP 37046 (KBEAUTY / OHADJA)
a partir de docs/DOCUMENT-COMPLET-RNCP.md.

- Titres H1/H2/H3 -> styles Heading + couleur prune #6B2D5C, filet or #9C7A3C
- Paragraphes, listes a puces, listes numerotees, blockquotes
- Tableaux markdown -> vrais tableaux Word
- Blocs mermaid -> insertion du PNG correspondant (docs/diagrammes/)
- Page de garde + pied de page avec numero de page
- Robuste : une image absente est ignoree sans planter
"""

import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Chemins
# ---------------------------------------------------------------------------
BASE = r"C:/Users/adria/Documents/Projets/PROJET/projet-miroirIoT"
SRC_MD = os.path.join(BASE, "docs", "DOCUMENT-COMPLET-RNCP.md")
DIAG_DIR = os.path.join(BASE, "docs", "diagrammes")
OUT_DOCX = os.path.join(BASE, "docs", "Dossier-KBEAUTY-RNCP.docx")

# ---------------------------------------------------------------------------
# Identite visuelle
# ---------------------------------------------------------------------------
PRUNE = RGBColor(0x6B, 0x2D, 0x5C)
OR_FILET = "9C7A3C"          # hex sans # pour XML
PRUNE_HEX = "6B2D5C"
BODY_FONT = "Calibri"
HEAD_FONT = "Cambria"

# Mapping mot-cle (contenu du bloc mermaid) -> images a inserer (legende)
# Les blocs mermaid presents dans le markdown : cas d'usage, sequence, deploiement.
MERMAID_RULES = [
    ("sequenceDiagram", [("04-sequence-seance.png", "Diagramme de sequence - workflow d'une seance")]),
    ("S'authentifier", [("03-cas-usage.png", "Diagramme de cas d'usage")]),
    ("Raspberry Pi 5 - Debian 12 kiosk", [
        ("05-deploiement.png", "Diagramme de deploiement"),
        ("07-architecture-composants.png", "Diagramme d'architecture des composants"),
        ("06-classes-services-device.png", "Diagramme de classes - services du device"),
    ]),
]

# Images inserees juste apres certains titres (texte du titre sans la numerotation markdown)
HEADER_IMAGES = {
    "4.3 Backlog reconstitue (extrait)": [("08-planning-gantt.png", "Planning previsionnel (Gantt)")],
    "5.3 Modelisation UML": [
        ("01-mcd.png", "Modele Conceptuel de Donnees (MCD) - 8 entites"),
        ("02-mld.png", "Modele Logique de Donnees (MLD)"),
    ],
}


# ---------------------------------------------------------------------------
# Helpers XML / style
# ---------------------------------------------------------------------------
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_bottom_border(paragraph, hex_color, size="12"):
    """Filet or sous un paragraphe (utilise pour les titres H1)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_page_number_field(paragraph):
    """Insere un champ PAGE dans un paragraphe (pied de page)."""
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldEnd)


# Tokenizer inline : **gras**, *italique*, `code`
INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(paragraph, text, base_bold=False, base_italic=False, color=None):
    if text is None:
        return
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        bold = base_bold
        italic = base_italic
        mono = False
        content = part
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            bold = True
            content = part[2:-2]
        elif part.startswith("*") and part.endswith("*") and len(part) >= 2:
            italic = True
            content = part[1:-1]
        elif part.startswith("`") and part.endswith("`") and len(part) >= 2:
            mono = True
            content = part[1:-1]
        run = paragraph.add_run(content)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        if color is not None:
            run.font.color.rgb = color


# ---------------------------------------------------------------------------
# Insertion d'image
# ---------------------------------------------------------------------------
def usable_width(document):
    sec = document.sections[0]
    return sec.page_width - sec.left_margin - sec.right_margin


def insert_image(document, filename, caption):
    path = os.path.join(DIAG_DIR, filename)
    if not os.path.isfile(path):
        return False  # robuste : on ignore l'image absente
    try:
        max_w = usable_width(document)
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=max_w)
    except Exception:
        return False
    if caption:
        cap = document.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = PRUNE
    return True


# ---------------------------------------------------------------------------
# Tableaux
# ---------------------------------------------------------------------------
def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_separator_row(line):
    cells = split_row(line)
    if not cells:
        return False
    return all(re.fullmatch(r":?-{2,}:?", c.strip()) for c in cells if c.strip() != "") and any(
        "-" in c for c in cells
    )


def add_table(document, rows):
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]
    ncol = len(header)
    table = document.add_table(rows=1, cols=ncol)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    has_header = any(h.strip() for h in header)
    hdr_cells = table.rows[0].cells
    for i in range(ncol):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        txt = header[i] if i < len(header) else ""
        add_runs(p, txt, base_bold=True, color=RGBColor(0xFF, 0xFF, 0xFF) if has_header else None)
        if has_header:
            set_cell_background(cell, PRUNE_HEX)

    for brow in body:
        cells = table.add_row().cells
        for i in range(ncol):
            cell = cells[i]
            cell.text = ""
            txt = brow[i] if i < len(brow) else ""
            add_runs(cell.paragraphs[0], txt)
    document.add_paragraph()  # espace apres tableau


# ---------------------------------------------------------------------------
# Page de garde
# ---------------------------------------------------------------------------
def build_cover(document):
    for _ in range(2):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Miroir connecte d'analyse capillaire")
    r.font.name = HEAD_FONT
    r.font.size = Pt(30)
    r.bold = True
    r.font.color.rgb = PRUNE

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KBEAUTY - service Bubble Hair Spa")
    r.font.name = HEAD_FONT
    r.font.size = Pt(18)
    r.font.color.rgb = PRUNE

    sep = document.add_paragraph()
    add_bottom_border(sep, OR_FILET, size="18")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dossier de soutenance")
    r.font.size = Pt(16)
    r.bold = True
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Titre RNCP 37046 - "Chef de projet en solutions logicielles pour l\'Internet des Objets"'
    )
    r.font.size = Pt(13)
    r.italic = True

    for _ in range(2):
        document.add_paragraph()

    infos = [
        ("Candidat", "Adriano Palamara - developpeur junior full-stack, specialise IoT"),
        ("Entreprise (alternance)", "OHADJA - SAS, programmation informatique (Paris 8e)"),
        ("Formation", "ACADENICE - Bachelor 3 (niveau 6) - certificateur FACILITYCERT"),
        ("Client", "KBEAUTY-COSMETICS - institut K-beauty (Nice)"),
        ("Date de soutenance", "2026"),
    ]
    for label, value in infos:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label + " : ")
        r.bold = True
        r.font.color.rgb = PRUNE
        p.add_run(value)

    document.add_page_break()


# ---------------------------------------------------------------------------
# Configuration des styles globaux
# ---------------------------------------------------------------------------
def setup_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)

    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        st = document.styles[name]
        st.font.name = HEAD_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = PRUNE

    # Pied de page : numero de page centre
    footer = document.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("KBEAUTY - RNCP 37046  |  page ").font.size = Pt(9)
    add_page_number_field(fp)


# ---------------------------------------------------------------------------
# Parsing markdown
# ---------------------------------------------------------------------------
def clean_header_text(text):
    # retire les ** eventuels pour le mapping
    return text.replace("**", "").strip()


def build_document():
    with open(SRC_MD, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    document = Document()
    setup_styles(document)
    build_cover(document)

    # On saute l'en-tete markdown (titre, table d'identite, note) deja traite
    # par la page de garde : on demarre au SOMMAIRE.
    start = 0
    for i, ln in enumerate(lines):
        if ln.strip() == "# SOMMAIRE":
            start = i
            break
    lines = lines[start:]

    n = len(lines)
    i = 0
    section_count = 0

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # --- Bloc de code / mermaid ---
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # ferme le fence
            content = "\n".join(block)
            if lang == "mermaid" or "mermaid" in lang:
                matched = False
                for keyword, imgs in MERMAID_RULES:
                    if keyword in content:
                        for fname, cap in imgs:
                            insert_image(document, fname, cap)
                        matched = True
                        break
                if not matched:
                    # bloc mermaid non reconnu : on ignore silencieusement
                    pass
            else:
                # code generique -> bloc monospace
                p = document.add_paragraph()
                r = p.add_run(content)
                r.font.name = "Consolas"
                r.font.size = Pt(9)
            continue

        # --- Tableaux ---
        if stripped.startswith("|") and (i + 1) < n and is_separator_row(lines[i + 1]):
            tbl = [lines[i], lines[i + 1]]
            i += 2
            while i < n and lines[i].strip().startswith("|"):
                tbl.append(lines[i])
                i += 1
            add_table(document, tbl)
            continue

        # --- Titres ---
        if stripped.startswith("# "):
            text = clean_header_text(stripped[2:])
            if section_count > 0:
                document.add_page_break()
            h = document.add_heading(level=1)
            add_runs(h, text)
            add_bottom_border(h, OR_FILET, size="12")
            section_count += 1
            i += 1
            continue
        if stripped.startswith("## "):
            text = clean_header_text(stripped[3:])
            h = document.add_heading(level=2)
            add_runs(h, text)
            i += 1
            if text in HEADER_IMAGES:
                for fname, cap in HEADER_IMAGES[text]:
                    insert_image(document, fname, cap)
            continue
        if stripped.startswith("### "):
            text = clean_header_text(stripped[4:])
            h = document.add_heading(level=3)
            add_runs(h, text)
            i += 1
            continue

        # --- Separateur horizontal ---
        if stripped == "---":
            i += 1
            continue

        # --- Blockquote ---
        if stripped.startswith(">"):
            qtext = stripped[1:].strip()
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_bottom_border(p, OR_FILET, size="6")
            add_runs(p, qtext, base_italic=True, color=PRUNE)
            i += 1
            continue

        # --- Listes a puces ---
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = document.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        # --- Listes numerotees ---
        m = re.match(r"^\d+\.\s+(.*)$", stripped)
        if m:
            p = document.add_paragraph(style="List Number")
            add_runs(p, m.group(1))
            i += 1
            continue

        # --- Ligne vide ---
        if stripped == "":
            i += 1
            continue

        # --- Paragraphe normal ---
        p = document.add_paragraph()
        add_runs(p, stripped)
        i += 1

    document.save(OUT_DOCX)
    return section_count


if __name__ == "__main__":
    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    count = build_document()
    size = os.path.getsize(OUT_DOCX)
    print("OK ->", OUT_DOCX)
    print("Taille (octets):", size)
    print("Sections H1:", count)
